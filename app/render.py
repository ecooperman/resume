"""
The actual rendering pipeline: resume-yaml-shaped data -> HTML / PDF / DOCX.

Split out of build.py (which is still the only thing that writes to
site/ - it just imports these functions now) so that app/admin.py can also
reach them, for rendering *tailored* resumes (from the time-management app's
resume-generation feature) into ad-hoc PDF/DOCX bytes without touching
site/ output at all. See app/admin.py's /api/render endpoint.

Every real person's resume content lives in the `people` table (models.Person)
- never a file, never in git (see app/crud.py). This module only ever
receives already-loaded data as a plain dict; it doesn't know or care where
that dict came from.
"""
from datetime import date
from pathlib import Path

from jinja2 import Environment, FileSystemLoader
from playwright.sync_api import sync_playwright
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).resolve().parent.parent  # repo root, one level above app/
TEMPLATES_DIR = ROOT / "templates"


def render_html(data, *, show_download_bar):
    env = Environment(loader=FileSystemLoader(TEMPLATES_DIR), autoescape=True)
    template = env.get_template("resume.html.j2")
    return template.render(show_download_bar=show_download_bar, **data)


def build_pdf(data, out_dir):
    # Render without the download bar -- it's a page-only UI element.
    html = render_html(data, show_download_bar=False)
    tmp = out_dir / "_pdf_render.html"
    tmp.write_text(html)
    try:
        with sync_playwright() as p:
            browser = p.chromium.launch()
            page = browser.new_page()
            page.goto(tmp.as_uri())
            out = out_dir / "resume.pdf"
            page.pdf(
                path=str(out),
                format="Letter",
                print_background=True,
                margin={"top": "0in", "bottom": "0in", "left": "0in", "right": "0in"},
            )
            browser.close()
        print(f"wrote {out}")
    finally:
        tmp.unlink(missing_ok=True)


# --- Cover letters ---
#
# Rendered from just basics (for the letterhead) + the generated body text
# (see time-management's app/resume_gen.py), not a full resume-shaped dict -
# there's no per-person cover_letter.yaml, this is always ad-hoc generated
# content passed straight through from the caller.


def _cover_letter_paragraphs(cover_letter_text):
    """Splits on blank lines into paragraphs, each itself a list of lines
    (split on single newlines within that paragraph) - a sign-off like
    "Sincerely,\\nEvan Cooperman" is a single block with no blank line in
    it, and naively dropping single newlines would silently collapse that
    onto one line. Callers render each line separately with an explicit
    break between them rather than joining on a bare "\\n", which HTML (and
    a single docx run) both ignore."""
    text = (cover_letter_text or "").strip()
    blocks = [b.strip() for b in text.split("\n\n") if b.strip()]
    if not blocks and text:
        blocks = [text]
    return [[line.strip() for line in block.splitlines() if line.strip()] for block in blocks]


def render_cover_letter_html(basics, cover_letter_text):
    env = Environment(loader=FileSystemLoader(TEMPLATES_DIR), autoescape=True)
    template = env.get_template("cover_letter.html.j2")
    return template.render(
        basics=basics,
        paragraphs=_cover_letter_paragraphs(cover_letter_text),
        today=date.today().strftime("%B %-d, %Y"),
    )


def build_cover_letter_pdf(basics, cover_letter_text, out_dir):
    html = render_cover_letter_html(basics, cover_letter_text)
    tmp = out_dir / "_cover_letter_render.html"
    tmp.write_text(html)
    try:
        with sync_playwright() as p:
            browser = p.chromium.launch()
            page = browser.new_page()
            page.goto(tmp.as_uri())
            out = out_dir / "cover_letter.pdf"
            page.pdf(
                path=str(out),
                format="Letter",
                print_background=True,
                # Margins come entirely from the template's @media print
                # padding (same split as resume.html.j2's build_pdf above) -
                # zero here so they aren't applied twice.
                margin={"top": "0in", "bottom": "0in", "left": "0in", "right": "0in"},
            )
            browser.close()
        print(f"wrote {out}")
    finally:
        tmp.unlink(missing_ok=True)


DEFAULT_FONT = "Calibri"
ACCENT = (0x24, 0x54, 0xFF)
MUTED = (0x5B, 0x62, 0x70)
INK = (0x1A, 0x1D, 0x23)


def build_docx(data, out_dir):
    basics = data["basics"]
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = DEFAULT_FONT
    style.font.size = Pt(10.5)
    for section in doc.sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    _run(doc.add_paragraph(), basics["name"], bold=True, size=22, color=INK)

    _run(doc.add_paragraph(), basics["title"], bold=True, size=12, color=ACCENT)

    contact_bits = [basics.get("location", ""), basics.get("email", ""), basics.get("phone", "")]
    for key in ("website", "linkedin", "github"):
        if basics.get(key):
            contact_bits.append(basics[key])
    _run(doc.add_paragraph(), " · ".join(b for b in contact_bits if b), size=9.5, color=MUTED)

    if basics.get("summary"):
        doc.add_paragraph()
        _run(doc.add_paragraph(), basics["summary"].strip(), size=10, color=MUTED)

    _heading(doc, "Experience")
    for job in data.get("experience", []):
        p = doc.add_paragraph()
        _run(p, job["title"], bold=True, size=11, color=INK)
        _run(p, f" - {job['company']}", size=10, color=INK)
        _run(p, f"\t{job.get('start', '')} – {job.get('end', '')}", size=9.5, color=MUTED)
        if job.get("location"):
            _run(doc.add_paragraph(), job["location"], size=9.5, color=MUTED)
        for bullet in job.get("bullets", []):
            _run(doc.add_paragraph(style="List Bullet"), bullet, size=10, color=INK)

    _heading(doc, "Education")
    for ed in data.get("education", []):
        p = doc.add_paragraph()
        title = f"{ed.get('degree', '')}"
        if ed.get("field"):
            title += f", {ed['field']}"
        _run(p, title, bold=True, size=11, color=INK)
        if ed.get("start") or ed.get("end"):
            _run(p, f"\t{ed.get('start', '')} – {ed.get('end', '')}", size=9.5, color=MUTED)
        _run(doc.add_paragraph(), ed.get("school", ""), size=9.5, color=MUTED)

    if data.get("certifications"):
        _heading(doc, "Certifications")
        for cert in data["certifications"]:
            p = doc.add_paragraph(style="List Bullet")
            _run(p, cert["name"], size=10, color=INK)
            if cert.get("issuer"):
                _run(p, f" - {cert['issuer']}", size=9.5, color=MUTED)

    _heading(doc, "Skills")
    for s in data.get("skills", []):
        p = doc.add_paragraph()
        _run(p, f"{s['category']}: ", bold=True, size=10, color=INK)
        _run(p, ", ".join(s.get("items", [])), size=10, color=INK)

    if data.get("projects"):
        _heading(doc, "Projects")
        for proj in data["projects"]:
            p = doc.add_paragraph()
            _run(p, proj["name"], bold=True, size=10.5, color=INK)
            if proj.get("link"):
                _run(p, f" - {proj['link']}", size=9.5, color=MUTED)
            _run(doc.add_paragraph(), proj.get("description", ""), size=10, color=INK)

    out = out_dir / "resume.docx"
    doc.save(out)
    print(f"wrote {out}")


def _heading(doc, text):
    doc.add_paragraph()
    p = doc.add_paragraph()
    _run(p, text.upper(), bold=True, size=10.5, color=MUTED)
    p.paragraph_format.space_after = Pt(4)


def _run(paragraph, text, *, bold=False, size=10.5, color=None):
    """Add a run with the font explicitly and consistently set (ascii + east
    Asian), instead of relying on style inheritance -- some renderers apply
    the document's default font inconsistently between runs otherwise."""
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = DEFAULT_FONT
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), DEFAULT_FONT)
    if color is not None:
        run.font.color.rgb = RGBColor(*color)
    return run


def build_cover_letter_docx(basics, cover_letter_text, out_dir):
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = DEFAULT_FONT
    style.font.size = Pt(11)
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)

    _run(doc.add_paragraph(), basics.get("name", ""), bold=True, size=16, color=INK)
    contact_bits = [basics.get(k, "") for k in ("email", "phone", "location")]
    _run(doc.add_paragraph(), " · ".join(b for b in contact_bits if b), size=9.5, color=MUTED)

    doc.add_paragraph()
    _run(doc.add_paragraph(), date.today().strftime("%B %-d, %Y"), size=10.5, color=MUTED)
    doc.add_paragraph()

    for lines in _cover_letter_paragraphs(cover_letter_text):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(10)
        last_run = None
        for line in lines:
            if last_run is not None:
                last_run.add_break()
            last_run = _run(p, line, size=11, color=INK)

    out = out_dir / "cover_letter.docx"
    doc.save(out)
    print(f"wrote {out}")
